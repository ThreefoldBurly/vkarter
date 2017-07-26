# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from fuzzywuzzy import fuzz

from vkarter_input import *
from vkarter_shared import *


# TODO: obsługa PLACEHOLDERÓW
# TODO: pomijanie adresów w Miejscu Pomiaru
# TODO: prawidłowa obsługa (pomijanie) cudzysłowów
class GeneratorNazwyPliku(object):
    """Generuje odpowiednio sformatowaną nazwę pliku i przekazuje obiektowi klasy Stanowisko"""
    # stałe klasy
    DWUZNAKI = (u'Ch', u'Cz', u'Dz', u'Dź', u'Dż', u'Rz', u'Sz')
    NIELEGALNE = u'/\?%*:|"<>'

    def __init__(self, nr_st, miejsce_pom, nazwa_st, nr_spr):
        super(GeneratorNazwyPliku, self).__init__()
        self.nr_st = nr_st
        self.miejsce_pom = miejsce_pom
        self.nazwa_st = nazwa_st
        self.nr_spr = (u"[spr_" + nr_spr + u"]").replace(u"/", u"-")

    def generujSkrotMiejscaPom(self):
        """Generuje skrót miejsca pomiaru dla nazwy pliku"""
        popr_miejsce = self.miejsce_pom.replace(u": ", u" ").replace(u"-", u" ").replace(u"   ", u" ").replace(u", ", u",").replace(u"\n", u",")

        podzial = popr_miejsce.split(u",")

        nowy_podzial = []
        for fraza in podzial:
            nowa_fraza = u""
            wyrazy = fraza.split()

            for wyraz in wyrazy:
                # obsługa nawiasów
                byl_nawias = False
                if wyraz[0] == u"(":
                    nowa_fraza += u"("
                    wyraz = wyraz[1:]
                if wyraz[-1] == u")":
                    byl_nawias = True
                    wyraz = wyraz[:-1]

                # obsługa liczb
                if wyraz.isdigit():
                    nowa_fraza += wyraz
                    if byl_nawias:
                        nowa_fraza += u")"
                    continue

                # obsługa liczb porządkowych (rzymskich)
                rzymskie = konwertujRzymskie(wyraz)
                if rzymskie.isdigit():
                    nowa_fraza += rzymskie
                    if byl_nawias:
                        nowa_fraza += u")"
                    continue

                # obsługa skrótowców
                if wyraz.isupper():
                    nowa_fraza += wyraz
                    if byl_nawias:
                        nowa_fraza += u")"
                    continue

                ile_znakow = 1
                if wyraz[0].isupper():
                    if len(wyrazy) == 1:
                        ile_znakow += 1
                    if wyraz[0:2] in self.DWUZNAKI:
                        ile_znakow += 1

                nowa_fraza += wyraz[:ile_znakow]
                if byl_nawias:
                    nowa_fraza += u")"

            nowy_podzial.append(nowa_fraza)

        return u",".join(nowy_podzial)

    def generujSkrotNazwySt(self):
        """
        Na podstawie znanego klasie słownika skrótów generuje skrótową wersję nazwy stanowiska do umieszczenia w nazwie pliku
        """
        slownik_skrotow = STALE_ZEWNETRZNE[u"SŁOWNIK_SKRÓTÓW"]
        skrot = self.nazwa_st

        for klucz in slownik_skrotow:
            if klucz in self.nazwa_st:
                skrot = skrot.replace(klucz, slownik_skrotow[klucz])

        return skrot

    # tylko legalne znaki w nazwie!
    def generujNazwePliku(self, suffix, prog=50):
        """
        Tworzy nazwę pliku (karty lub rejestru w zależności od suffiksu), skracając nazwę stanowiska, jeśli ilość znaków w wygenerowanej nazwie przekroczy podany próg
        """
        skrot_miejsca = self.generujSkrotMiejscaPom()
        wstepna_dlugosc = 3 + 1 + len(skrot_miejsca) + 1 + len(self.nazwa_st) + 1 + len(self.nr_spr) + len(suffix)

        nazwa_pliku = []
        nazwa_pliku.append(self.nr_st)
        nazwa_pliku.append(u"_")
        nazwa_pliku.append(skrot_miejsca)
        nazwa_pliku.append(u" ")
        if wstepna_dlugosc >= prog:
            skrot_nazwy = self.generujSkrotNazwySt()

            # obsługa za długich nazw
            zadlugosc = len(skrot_miejsca) + len(skrot_nazwy)
            if zadlugosc > 225:
                skrot_nazwy = skrot_nazwy[:160]

            nazwa_pliku.append(skrot_nazwy)
        else:
            nazwa_pliku.append(self.nazwa_st)
        nazwa_pliku.append(u" ")
        nazwa_pliku.append(self.nr_spr)
        nazwa_pliku.append(suffix)
        nazwa_pliku.append(u".docx")

        nazwa_pliku = u"".join(nazwa_pliku)

        # test legalności znaków
        for znak in nazwa_pliku:
            if znak in self.NIELEGALNE:
                nazwa_pliku = nazwa_pliku.replace(znak, u"-")

        # test długości
        if len(nazwa_pliku) > 229:
            print u"Długość nazwy pliku: %d" % len(nazwa_pliku)
            nazwa_pliku = nazwa_pliku[:229]
            print u"Długość nazwy pliku po obcięciu: %d" % len(nazwa_pliku)

        return nazwa_pliku


# TODO: dodać pole zachowujące wygenerowaną nazwę pliku (generowana nazwa nie powinna wtedy zawierać numeru, a numeracja powinna zależeć od nazwy a konkretnie od skrótu miejsca pomiaru). Być może dobrze byłoby stworzyć osobną klasę dla nazwy pliku albo stworzyc osobną klasę dla karty i przeniesc to polę tam
class Stanowisko(object):
    """
    Abstrakcyjna reprezentacja stanowiska - realizowana jako Karta Czynników Szkodliwych lub Rejestr Czynników Szkodliwych
    """

    def __init__(self, czynniki, numer):
        super(Stanowisko, self).__init__()
        self.czynniki = czynniki  # lista
        self.numer = numer
        self.karta = self.przygotujKarte()

    @staticmethod
    def przygotujKarte():
        """Inicjuje plik .docx dla karty"""
        karta = Document()

        # ustawienia strony
        sekcja = karta.sections[0]

        sekcja.page_width = Cm(21.0)
        sekcja.page_height = Cm(29.7)

        sekcja.top_margin = Cm(1.0)
        sekcja.bottom_margin = Cm(1.0)
        sekcja.left_margin = Cm(2.0)
        sekcja.right_margin = Cm(2.0)

        # przygotowanie nowych styli
        style = karta.styles

        # style nagłówka
        styl_an = style.add_style('akapit-naglowka', WD_STYLE_TYPE.PARAGRAPH)
        styl_an.base_style = style['Normal']
        styl_an.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styl_an.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        styl_an.paragraph_format.line_spacing = 1.0
        styl_an.paragraph_format.space_before = Pt(3)
        styl_an.paragraph_format.space_after = Pt(3)
        styl_an.paragraph_format.widow_control = True

        styl_zn = style.add_style('znaki-naglowka', WD_STYLE_TYPE.CHARACTER)
        styl_zn.base_style = style['Default Paragraph Font']
        styl_zn.font.name = 'Times New Roman'
        styl_zn.font.size = Pt(10)
        styl_zn.font.bold = True

        # style tabelki
        styl_tab = style.add_style('tabelka', WD_STYLE_TYPE.TABLE)
        styl_tab.base_style = style['Table Grid']
        styl_tab.paragraph_format.space_before = Pt(1.5)
        styl_tab.paragraph_format.space_after = Pt(1.5)
        styl_tab.paragraph_format.keep_together = True

        styl_zt = style.add_style('znaki-tabelki', WD_STYLE_TYPE.CHARACTER)
        styl_zt.base_style = style['Default Paragraph Font']
        styl_zt.font.name = 'Times New Roman'
        styl_zt.font.size = Pt(10)

        # style stopki (wykorzystywane tylko pod czynnikami chemicznymi)
        styl_as = style.add_style('akapit-stopki', WD_STYLE_TYPE.PARAGRAPH)
        styl_as.base_style = style['Normal']
        styl_as.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        styl_as.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        styl_as.paragraph_format.line_spacing = 1.0
        styl_as.paragraph_format.space_before = Pt(0)
        styl_as.paragraph_format.space_after = Pt(0)
        styl_as.paragraph_format.widow_control = True

        styl_zs = style.add_style('znaki-stopki', WD_STYLE_TYPE.CHARACTER)
        styl_zs.base_style = style['Default Paragraph Font']
        styl_zs.font.name = 'Times New Roman'
        styl_zs.font.size = Pt(8)

        tytul1 = karta.add_paragraph(style='akapit-naglowka')
        tytul1.add_run(u'KARTA BADAŃ I POMIARÓW CZYNNIKÓW SZKODLIWYCH', style='znaki-naglowka')
        # interlinia
        tytul2 = karta.add_paragraph(style='akapit-stopki')
        tytul2.add_run(u"", style='znaki-naglowka')

        return karta

    def utworzKarte(self):
        """Tworzy Kartę Czynników Szkodliwych dla tego stanowiska"""
        suffix = u"-karta"

        for czynnik in self.czynniki:
            czynnik.rysujTabelke(self.karta)

        # dane dla generatora pobierane są z pierwszego Czynnika
        generator_nazwy = GeneratorNazwyPliku(self.numer, self.czynniki[0].miejsce, self.czynniki[0].nazwa_st, self.czynniki[0].nr_spr)
        print u"Tworzę kartę!"
        self.karta.save(generator_nazwy.generujNazwePliku(suffix))

    def utworzRejestr(self):
        """Tworzy Rejestr Czynników Szkodliwych dla tego stanowiska"""
        pass

# class GeneratorCzynnikow(object):
#     """Generuje listy obiektów klasy Czynnik dla MonteraStanowisk z Pomiarów otrzymanych od ParseraDRK przy pomocy parserów odpowiednich typów."""
#     def __init__(self, pomiary):
#         super(GeneratorCzynnikow, self).__init__()
#         self.pomiary = pomiary # listy wszystkich zebranych Pomiarów uporządkowane wg typów w krotce otrzymywanej od obiektu klasy ParserDRK

#     def generujHalasy(self):
#         """Z obiektów zebranych od parserów generuje listę hałasów"""
#         halasy = [] # lista obiektów klasy Halas

#         for pomiar in self.pomiary.halas:
#             parser = ParserHalasu(pomiar)
#             halasy.append(parser.podajHalas())

#         print u"Długość listy 'halasy' na koniec metody generujHalasy(): %d" % len(halasy)

#         return halasy


#     def stworzPylyChemie(self):
#         pass

class MonterStanowisk(object):
    """
    Z Pomiarów otrzymanych od ParseraDRK generuje listy Czynników, z których z kolei, kojarząc je ze sobą właściwie na ile to możliwe, montuje stanowiska
    """

    def __init__(self, pomiary):
        super(MonterStanowisk, self).__init__()
        self.pomiary = pomiary

        self.halasy = self.stworzHalasy()  # lista obiektów
        # self.drgania = drgania
        self.pyly_chemie = self.stworzPylyChemie()  # lista list obiektów!

        self.stanowiska = []

        # print u"Długość pól MonteraStanowisk po inicjalizacji. Halas: %d" % len(self.halasy)

    def stworzHalasy(self):
        """Generuje Hałasy z Pomiarów odpowedniego typu"""
        halasy = []

        for pomiar in self.pomiary.halas:
            ph = ParserHalasu(pomiar)
            halasy.append(ph.podajHalas())

        print u"Długość listy 'halasy' na koniec metody stworzHalasy(): %d" % len(halasy)

        return halasy

    def stworzDrgania(self):
        pass

    def stworzPylyChemie(self):
        """Generuje Pyły/Chemie z Pomiarów odpowiedniego typu"""
        pyly_chemie = []

        for pomiar in self.pomiary.pyly_chemia:
            ppch = ParserPylowChemii(pomiar)
            pyly_chemie.append(ppch.podajCzynniki())

        print u"Długość listy 'pyly_chemie' na koniec metody stworzPylyChemie(): %d" % len(pyly_chemie)

        return pyly_chemie

    # TODO: dodać wybór pomiędzy użyciem formatu 2-cyfrowego i 3-cyfrowego w zależności od ilość wszystkich stanowisk i dodać obsługę tego w montujStanowiska() (to raczej oznacza dodawanie nr. do stanowiska poza konstruktorem)
    @staticmethod
    def generujNrStanowiska(numer):
        """Generuje nr stanowiska, który pojawia się w nazwie pliku wynikowego"""
        nr_st = numer + 1
        if nr_st >= 1 and nr_st <= 9:
            return "00" + str(nr_st)
        elif nr_st > 9 and nr_st < 100:
            return "0" + str(nr_st)
        else:
            return str(nr_st)

    @staticmethod
    def czySaPodobne(opis1, opis2, prog=95):
        """Porównuje podobieństwo opisów czynności, wykorzystując bibliotekę FuzzyWuzzy"""
        fuzzyratio = fuzz.token_set_ratio(opis1, opis2)
        if fuzzyratio >= prog:
            return True
        else:
            return False

    # TODO: dodać obsługę skojarzeń w obrębie jednego typu Pomiarów (zdarzało się tak w niektórych sprawozdaniach - te same stanowiska w różnych sprawozdaniach tego samego typu (osobno pył i oleje mineralne)). Dodać fuzzy matching do nazw stanowisk (jeśli to możliwe) - były takie sprawozdania, w których część nazw się nie skojarzyło bo wersja nazwy u mnie miała spacje w oznaczeniu maszyny a wersja tej samej nazwy u TT tej spacji nie miałą
    # TODO: stworzyć bardziej zaawansowane uporządkowywanie stanowisk (niż tylko: najpierw wszystkie stanowiska wg kolejności hałasu, potem reszta). Należałoby zczytać porządek wydziałów w hałasie (najlepiej na podstawie skrótów - ale to wymaga obudowania nazwy pliku klasą i wydzielenie skrótu miejsca pomiaru jako pola tej klasy) i zbudować nową listę wg tego porządku z wszystkich stworzonych stanowisk
    # TODO: dodać obsługę niepełnych Czynników (z brakującymi (jednocześnie) miejscem pomiaru i nazwą stanowiska) - powstałych z Pomiarów, będących tylko kontynuacją Pomiarów poprzedzających je w pliku DRK. Takie Czynniki byłyby automatycznie dołączane do stanowiska utworzonego dla poprzedniego Czynnika z listy Czynników i miałyby od razu uzupełniane dane o miejscu pomiaru i nazwie stanowiska.
    def montujStanowiska(self):
        """
        Kojarzy Czynniki przynależne do tego samego stanowiska i tworzy z nich obiekty klasy Stanowisko
        """
        dek_halasow = self.halasy[:]
        dek_pylow_chemii = self.pyly_chemie[:]
        print "Ilosc elementow w deku chemi: %d" % len(dek_pylow_chemii)

        licznik = 0
        for i in xrange(len(dek_halasow)):
            print "========= ITERACJA #%d =========" % (i + 1)
            print "nr stanowiska: %s" % self.generujNrStanowiska(i)
            czynniki_dla_st = []
            halas = dek_halasow[i]
            nazwa = halas.nazwa_st

            czynnosci_halasu = u", ".join(halas.czynnosci)

            # UWAGA! zmienna 'pylochem' jest tutaj listą Czynników
            # z listy pyłów/chemii wybieramy te, w których zgadza się nazwa stanowiska
            wybrane_pylochemy = []
            for pylochem in dek_pylow_chemii:
                n1 = nazwa
                n2 = pylochem[0].nazwa_st
                print "Token set ratio dla nazw: *** %d ***" % fuzz.token_set_ratio(n1, n2)
                # zamieniłem test równości nazw stanowiska na fuzzy matching (z progiem 97). Taki test wybiera pary nazw które różnią się 1-2 spacjami np.: "magazynier-operator", "magazynier - operator" czy "operator kuźniarki 88 RMC", "operator kuźniarki 88RMC" ale odrzuca te różniące się np. numeracją: "szlifierz VI", "szlifierz V" czy "szlifierz 1", "szlifierz 2"
                if self.czySaPodobne(n1, n2, 97):
                    wybrane_pylochemy.append(pylochem)

            # jeśli był tylko jeden - sprawdzmy czy opisy czynności się zgadzają. Przyjąłem fuzzy ratio: 95 (token_set)
            if len(wybrane_pylochemy) == 1:
                print "Pierwsza galaz"
                # opisy czynnosci w halasie trzeba najpierw scalic
                s1 = czynnosci_halasu
                s2 = wybrane_pylochemy[0][0].czynnosci

                print "Token set ratio w pierwszej galezi: *** %d ***" % fuzz.token_set_ratio(s1, s2)

                if self.czySaPodobne(s1, s2):
                    print "PIERWSZY STRING"
                    print s1
                    print "DRUGI STRING"
                    print s2
                    czynniki_dla_st.append(halas)
                    for czynnik in wybrane_pylochemy[0]:
                        czynniki_dla_st.append(czynnik)

                    self.stanowiska.append(Stanowisko(czynniki_dla_st, self.generujNrStanowiska(i)))
                    # usuwamy z deku pyłów/chemmi wykorzystany element
                    dek_pylow_chemii.remove(wybrane_pylochemy[0])
                    print "Usuwam element. Liczba elementów: %d" % len(dek_pylow_chemii)

                else:
                    # te same nazwy ale inne opisy czynnosci - jednak rózne stanowiska, więc tylko dodajemy tylko halas
                    czynniki_dla_st.append(halas)
                    self.stanowiska.append(Stanowisko(czynniki_dla_st, self.generujNrStanowiska(i)))

                licznik += 1

            elif len(wybrane_pylochemy) > 1:
                print "Druga galaz"
                bylo_dopasowanie = False

                for pylochem in wybrane_pylochemy:
                    s1 = czynnosci_halasu
                    s2 = pylochem[0].czynnosci

                    print "Token set ratio w drugiej galezi: *** %d ***" % fuzz.token_set_ratio(s1, s2)

                    if self.czySaPodobne(s1, s2):
                        print "PIERWSZY STRING"
                        print s1
                        print "DRUGI STRING"
                        print s2
                        czynniki_dla_st.append(halas)
                        for czynnik in pylochem:
                            czynniki_dla_st.append(czynnik)

                        self.stanowiska.append(Stanowisko(czynniki_dla_st, self.generujNrStanowiska(i)))
                        dek_pylow_chemii.remove(pylochem)
                        print "Usuwam element. Liczba elementów: %d" % len(dek_pylow_chemii)
                        bylo_dopasowanie = True
                        break

                if not bylo_dopasowanie:
                    czynniki_dla_st.append(halas)
                    self.stanowiska.append(Stanowisko(czynniki_dla_st, self.generujNrStanowiska(i)))

                licznik += 1

            else:
                print "Trzecia galaz"
                czynniki_dla_st.append(halas)
                self.stanowiska.append(Stanowisko(czynniki_dla_st, self.generujNrStanowiska(i)))

                licznik += 1

        # resztę stanowisk należy utworzyć z tego co zostału z deku pylow/chemii
        for j in xrange(len(dek_pylow_chemii)):
            pylochem = dek_pylow_chemii[j]
            self.stanowiska.append(Stanowisko(pylochem, self.generujNrStanowiska(licznik + j)))


# TODO: przenieść większość wywoływanych tu metod do metody __init__ poszczególnych klas (analogicznie do przygotujKarte() klasy Stanowisko)
def main():
    """Funkcja wykonująca skrypt"""
    sparsujPlikiVKA()
    p = ParserDRK()
    ms = MonterStanowisk(p.pomiary)
    ms.montujStanowiska()

    print u"Długość listy 'stanowiska' MonteraStanowisk po wywołaniu montujStanowiska(): %d" % len(ms.stanowiska)

    for stanowisko in ms.stanowiska:
        stanowisko.utworzKarte()

if __name__ == "__main__":
    main()
