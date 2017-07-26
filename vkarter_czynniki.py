# -*- coding: utf-8 -*-
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Czynnik(object):
    """
    Abstrakcyjna rerezentacja czynnika - realizowana jako tabelka w Karcie Czynników Szkodliwych
    """
    # tę stałą ze względu na 6. element należy przesłonić w klasie Pylochem
    TEKSTY_METRYKA = [u"Nazwa czynnika", u"Data pomiaru", u"Miejsce pomiaru", u"Wykonujący pomiar", u"Metoda pomiaru", u"Wynik pomiaru", u"Interpretacja wyniku", u"Stanowisko pracy"]

    def __init__(self, nazwa, nr_spr, miejsce, nazwa_st, data, ndw, wdo, krotnosc, czynnosci, wykonawca, metoda, ilosc_wierszy_tab=8):
        super(Czynnik, self).__init__()
        self.nazwa = nazwa
        self.nr_spr = nr_spr
        self.miejsce = miejsce
        self.nazwa_st = nazwa_st
        self.data = data

        self.ndw = ndw  # najwyższa dopuszczalna wartość (NDS lub NDN)
        self.wdo = wdo  # wartość do oceny - poziom ekspozycji dla Hałasu i Drgań, wskaźnik narażenia dla Pyłów i Chemii
        self.krotnosc = krotnosc
        self.czynnosci = czynnosci  # opisy czynności (w Hałasie jest to lista stringów, w PyłachChemii jest to pojedyńczy string)

        self.wykonawca = wykonawca
        self.metoda = metoda

        self.ilosc_wierszy_tab = ilosc_wierszy_tab

    def generujLP(self):
        """Generuje listę liczb porządkowych dla pierwszej kolumny tabelki danego czynnika"""
        l_p = []
        for i in xrange(self.ilosc_wierszy_tab):
            l_p.append(u"%d." % (i + 1))
        return l_p

    # metoda przesłaniana w klasach potomnych
    def rysujTabelke(self, karta):
        """Inicjuje tabelkę dla klasy potomnej i wypełnia 2 pierwsze kolumny"""
        naglowek = karta.add_paragraph(style='akapit-naglowka')

        tabelka = karta.add_table(rows=self.ilosc_wierszy_tab, cols=3, style='tabelka')
        tabelka.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tabelka.autofit = False

        tabelka.columns[0].width = Pt(27.0)
        tabelka.columns[1].width = Pt(103.5)
        tabelka.columns[2].width = Pt(351.5)

        # wypełnienie lewej kolumny
        for l_p, komorka in zip(self.generujLP(), tabelka.columns[0].cells):
            komorka.paragraphs[0].add_run(l_p, style='znaki-tabelki')
            komorka.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # wypełnianie środkowej kolumny ze względu na inną wersję stałej TEKSTY_METRYKA w pyłochemach przeniesione do klas potomnych

        # wypełnienie wierszy #1-5 prawej kolumny
        # wypełnienie wiersza #1 (NAZWA CZYNNIKA)
        k1 = tabelka.columns[2].cells[0]
        k1.paragraphs[0].add_run(self.nazwa, style='znaki-tabelki')
        k1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wiersza #2 (DATA POMIARU)
        k2 = tabelka.columns[2].cells[1]
        k2.paragraphs[0].add_run(self.data, style='znaki-tabelki')
        k2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wiersza #3 (MIEJSCE POMIARU)
        k3 = tabelka.columns[2].cells[2]
        linie_miejsca = self.miejsce.split(u'\n')

        for i in xrange(len(linie_miejsca)):
            linia = linie_miejsca[i]
            k3.paragraphs[i].add_run(linia, style='znaki-tabelki')
            k3.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.LEFT
            if len(linie_miejsca) - i > 1:
                k3.add_paragraph()

        # wypełnienie wiersza #4 (WYKONAWCA POMIARU)
        k4 = tabelka.columns[2].cells[3]
        k4.paragraphs[0].add_run(self.wykonawca, style='znaki-tabelki')
        k4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wiersza #5 (METODA POMIARU)
        k5 = tabelka.columns[2].cells[4]
        k5.paragraphs[0].add_run(self.metoda, style='znaki-tabelki')
        k5.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wiersza #8 (STANOWISKO PRACY)
        k8 = tabelka.columns[2].cells[7]
        k8.paragraphs[0].add_run(self.nazwa_st, style='znaki-tabelki')
        k8.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        return (naglowek, tabelka)


class Halas(Czynnik):
    """Podtyp Czynnika I stopnia"""
    # stałe klasy
    TEKSTY_NAGLOWEK = u"CZYNNIK FIZYCZNY"

    # uwaga na spacje na końcu stringów
    TEKSTY_WYNIK = (u"Poziom ekspozycji na hałas [dB]: ", u"Maksymalny poziom A [dB]: ", u"Szczytowy poziom C [dB]: ")

    TEKSTY_WARTOSCI_DOP = (u"Wartość dopuszczalna poziomu ekspozycji na hałas [dB]: ", u"Wartość dopuszczalna maksymalnego poziomu dźwięku A [dB]: ", u"Wartość dopuszczalna szczytowego poziomu dźwięku C [dB]: ")

    TEKSTY_OCENA = (u"Nie stwierdzono przekroczenia dopuszczalnych wartości parametrów hałasu.", u"Stwierdzono przekroczenie wartości dopuszczalnej poziomu ekspozycji na hałas.", u"Stwierdzono przekroczenie wartości dopuszczalnej maksymalnego poziomu dźwięku A.", u"Stwierdzono przekroczenie wartości dopuszczalnej szczytowego poziomu dźwięku C.", u"Stwierdzono przekroczenie dopuszczalnych wartości parametrów hałasu.")

    TEKSTY_KROTNOSC = (u"Krotność dopuszczalnego poziomu ekspozycji: ", u"Krotność dopuszczalnego maksymalnego poziomu dźwięku A: ", u"Krotność dopuszczalnego szczytowego poziomu dźwięku C: ")

    def __init__(self, nr_spr, miejsce, nazwa_st, data, ndp_ekspoz, poz_ekspoz, krotnosc, czynnosci, wykonawca, metoda, maks_A, szczyt_C, ndp_maks_A, ndp_szczyt_C):
        super(Halas, self).__init__(u"Hałas", nr_spr, miejsce, nazwa_st, data, ndp_ekspoz, poz_ekspoz, krotnosc, czynnosci, wykonawca, metoda)
        self.maks_A = maks_A
        self.szczyt_C = szczyt_C

        self.ndp_maks_A = ndp_maks_A
        self.ndp_szczyt_C = ndp_szczyt_C

        self.krotnosc_maks_A = None
        self.krotnosc_szczyt_C = None
        self.ustalKrotnosciAiC()

    def ustalKrotnosciAiC(self):
        """
        Sprawdza czy maks_A i szczyt_C przekraczają wartości dopuszczalne. Jeśli tak, to oblicza właściwe krotności
        """
        if int(self.maks_A) > int(self.ndp_maks_A):
            self.krotnosc_maks_A = str(round((10 ** (float(int(self.maks_A) - int(self.ndp_maks_A)) / 20)), 2))

        if int(self.szczyt_C) > int(self.ndp_szczyt_C):
            self.krotnosc_szczyt_C = str(round((10 ** (float(int(self.szczyt_C) - int(self.ndp_szczyt_C)) / 20)), 2))

    def rysujTabelke(self, karta):
        """
        Rysuje tabelkę tego czynnika. Cześć rysowania oddaje do tej samej metody w klasie rodzica
        """
        # 2 pierwsze kolumny oraz część trzeciej wypełnia rodzic
        naglowek, tabelka = super(Halas, self).rysujTabelke(karta)
        naglowek.add_run(self.TEKSTY_NAGLOWEK, style='znaki-naglowka')

        # wypełnienie środkowej kolumny
        for x in xrange(len(tabelka.columns[1].cells)):
            komorka = tabelka.columns[1].cells[x]
            komorka.paragraphs[0].add_run(self.TEKSTY_METRYKA[x], style='znaki-tabelki')
            komorka.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wierszy #6-8 prawej kolumny
        # wypełnienie wiersza #6 (WYNIK POMIARU)
        k6 = tabelka.columns[2].cells[5]
        k6.paragraphs[0].add_run(self.TEKSTY_WYNIK[0] + self.wdo, style='znaki-tabelki')
        k6.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k6.add_paragraph()
        k6.paragraphs[1].add_run(self.TEKSTY_WYNIK[1] + self.maks_A, style='znaki-tabelki')
        k6.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k6.add_paragraph()
        k6.paragraphs[2].add_run(self.TEKSTY_WYNIK[2] + self.szczyt_C, style='znaki-tabelki')
        k6.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wiersza #7 (INTERPRETACJA WYNIKU)
        # (wartości dopuszczalne)
        k7 = tabelka.columns[2].cells[6]
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP[0] + self.ndw, style='znaki-tabelki')
        k7.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k7.add_paragraph()
        k7.paragraphs[1].add_run(self.TEKSTY_WARTOSCI_DOP[1] + self.ndp_maks_A, style='znaki-tabelki')
        k7.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k7.add_paragraph()
        k7.paragraphs[2].add_run(self.TEKSTY_WARTOSCI_DOP[2] + self.ndp_szczyt_C, style='znaki-tabelki')
        k7.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k7.paragraphs[2].paragraph_format.space_after = Pt(3)

        # obsługa oceny (chyba prościej było zrobić to na krotnościach)
        k7.add_paragraph()
        dodane_akapity = 0
        if float(self.wdo) <= float(self.ndw) and int(self.maks_A) <= int(self.ndp_maks_A) and int(self.szczyt_C) <= int(self.ndp_szczyt_C):
            k7.paragraphs[3].add_run(self.TEKSTY_OCENA[0], style='znaki-tabelki')
            k7.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif float(self.wdo) > float(self.ndw) and int(self.maks_A) > int(self.ndp_maks_A) and int(self.szczyt_C) > int(self.ndp_szczyt_C):
            k7.paragraphs[3].add_run(self.TEKSTY_OCENA[4], style='znaki-tabelki')
            k7.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            k7.paragraphs[3].add_run(self.TEKSTY_OCENA[1], style='znaki-tabelki')
            if int(self.maks_A) > int(self.ndp_maks_A):
                k7.add_paragraph()
                dodane_akapity += 1
                k7.paragraphs[3 + dodane_akapity].add_run(self.TEKSTY_OCENA[2], style='znaki-tabelki')
                k7.paragraphs[3 + dodane_akapity].alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif int(self.szczyt_C) > int(self.ndp_szczyt_C):
                k7.add_paragraph()
                dodane_akapity += 1
                k7.paragraphs[3 + dodane_akapity].add_run(self.TEKSTY_OCENA[3], style='znaki-tabelki')
                k7.paragraphs[3 + dodane_akapity].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # obsługa krotności
        k7.add_paragraph()
        k7.paragraphs[4 + dodane_akapity].add_run(self.TEKSTY_KROTNOSC[0] + self.krotnosc, style='znaki-tabelki')
        k7.paragraphs[4 + dodane_akapity].alignment = WD_ALIGN_PARAGRAPH.LEFT

        if self.krotnosc_maks_A is not None:
            k7.add_paragraph()
            dodane_akapity += 1
            k7.paragraphs[4 + dodane_akapity].add_run(self.TEKSTY_KROTNOSC[1] + self.krotnosc_maks_A, style='znaki-tabelki')
            k7.paragraphs[4 + dodane_akapity].alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif self.krotnosc_szczyt_C is not None:
            k7.add_paragraph()
            dodane_akapity += 1
            k7.paragraphs[4 + dodane_akapity].add_run(self.TEKSTY_KROTNOSC[2] + self.krotnosc_szczyt_C, style='znaki-tabelki')
            k7.paragraphs[4 + dodane_akapity].alignment = WD_ALIGN_PARAGRAPH.LEFT

        k7.paragraphs[4 + dodane_akapity].paragraph_format.space_before = Pt(3)
        k7.paragraphs[4 + dodane_akapity].paragraph_format.space_after = Pt(1)

        # interlinia
        karta.add_paragraph(style='akapit-stopki')


class Drgania(Czynnik):
    """Podtyp Czynnika I stopnia"""

    def __init__(self, nazwa, nr_spr, miejsce, nazwa_st, data, ndp_ekspoz, poz_ekspoz, krotnosc, czynnosci, wykonawca, metoda, numer, zrodlo, punkt_pom):
        super(Drgania, self).__init__(nazwa, nr_spr, miejsce, nazwa_st, data, ndp_ekspoz, poz_ekspoz, krotnosc, czynnosci, wykonawca, metoda)
        self.zrodlo = zrodlo
        self.punkt_pom = punkt_pom


class DrganiaOgolne(Drgania):
    """Podtyp Czynnika II stopnia reprezentujący drgania ogólne"""

    def __init__(self, nazwa, nr_spr, miejsce, nazwa_st, data, ndp_ekspoz, poz_ekspoz, krotnosc, czynnosci, wykonawca, metoda, numer, zrodlo, punkt_pom):
        super(DrganiaOgolne, self).__init__(nazwa, nr_spr, miejsce, nazwa_st, data, ndp_ekspoz, poz_ekspoz, krotnosc, czynnosci, wykonawca, metoda, zrodlo, punkt_pom)


class DrganiaMiejscowe(Drgania):
    """
    Podtyp Czynnika II stopnia reprezentujący drgania miejscowe z 1 lub 2 punktami pomiarowymi
    """

    def __init__(self, nazwa, nr_spr, miejsce, nazwa_st, data, ndp_ekspoz, poz_ekspoz, krotnosc, czynnosci, wykonawca, metoda, numer, zrodlo, punkt_pom, poz_ekspoz_dr):
        super(DrganiaMiejscowe, self).__init__(nazwa, nr_spr, miejsce, nazwa_st, data, ndp_ekspoz, poz_ekspoz, krotnosc, czynnosci, wykonawca, metoda, zrodlo, punkt_pom)
        self.poz_ekspoz_dr = poz_ekspoz_dr  # poziom ekspozycji dla drugiej ręki


class Pylochem(Czynnik):
    """
    Podtyp Czynnika I stopnia reprezentujący pyły bez krzemionki (pyły drewna) oraz chemię bez chwilówek (dozymetryczną i stacjonarną)
    """
    # stałe klasy
    TEKSTY_NAGLOWEK = (u"PYŁ", u"CZYNNIK CHEMICZNY")

    TEKSTY_METRYKA = [u"Nazwa czynnika", u"Data pomiaru", u"Miejsce pomiaru", u"Wykonujący pomiar", u"Metoda pomiaru", (u"Wynik pomiaru", u"- ocena NDS", u"- ocena NDSP*)"), u"Interpretacja wyniku", u"Stanowisko pracy", u"Uwagi"]
    # uwaga na spacje na końcu stringów
    # '3' musi być w superscripcie (Font object, superscript property, value: True) - dlatego tekst podzielony na 3 części
    TEKSTY_WYNIK = (u"Wskaźnik narażenia [mg/m", u"3", u"]: ")

    TEKSTY_WARTOSCI_DOP = (u"Wartość dopuszczalna wskaźnika narażenia [mg/m", u"3", u"]: ")

    TEKSTY_OCENA = (u"Wskaźnik narażenia nie przekracza wartości dopuszczalnej.", u"Wskaźnik narażenia przekracza wartość dopuszczalną.")

    TEKSTY_KROTNOSC = u"Krotność wartości dopuszczalnej wskaźnika narażenia: "

    TEKSTY_P_O = (u"Stężenie poniżej oznaczalności – poniżej ", u" [mg/m", u"3", u"]")

    TESKTY_STOPKA = u"*) – W przypadku pomiarów ciągłych – wartość maksymalna"

    def __init__(self, nazwa, nr_spr, miejsce, nazwa_st, data, nds, wskaznik_nar, krotnosc, czynnosci, wykonawca, metoda, numer, p_o):
        super(Pylochem, self).__init__(nazwa, nr_spr, miejsce, nazwa_st, data, nds, wskaznik_nar, krotnosc, czynnosci, wykonawca, metoda)
        self.numer = numer  # liczba porządkowa decydująca o kolejności w obrębie tego samego stanowiska
        self.p_o = p_o  # p.o. - poniżej oznaczalności
        self.ustalIloscWierszyTabeli()

    def ustalIloscWierszyTabeli(self):
        """Ustala ilość wierszy tabeli w zależności od tego, czy pole p_o w tym czynniku jest czy nie jest obiektem 'None'"""
        if self.p_o is not None:
            self.ilosc_wierszy_tab = 9

    # TODO: space_before i space_after nie działają prawidłowo (wina python-docx?)
    # UWAGA! Ta metoda w przeciwieństwie do metody w HAŁASIE zwraca wartość
    def rysujTabelke(self, karta):
        """Rysuje tabelkę tego czynnika. Cześć rysowania oddaje do tej samej metody w klasie rodzica"""
        # metody podręczna
        def sprawdzCzyPyl(nazwa_czynnika):
            """Sprawdza czy ten czynnik jest pyłem"""
            if u"pył" in self.nazwa.lower():
                return True
            else:
                return False

        # 2 pierwsze kolumny oraz część trzeciej wypełnia rodzic
        naglowek, tabelka = super(Pylochem, self).rysujTabelke(karta)
        jestem_pylem = sprawdzCzyPyl(self.nazwa)

        # nagłówek tabelki
        if jestem_pylem:
            naglowek.add_run(self.TEKSTY_NAGLOWEK[0], style='znaki-naglowka')
        else:
            naglowek.add_run(self.TEKSTY_NAGLOWEK[1], style='znaki-naglowka')

        # wypełnienie środkowej kolumny
        for x in xrange(len(tabelka.columns[1].cells)):
            komorka = tabelka.columns[1].cells[x]
            # obsługa 6tego wiersza ('Wynik pomiaru')
            if x == 5:
                komorka.paragraphs[0].add_run(self.TEKSTY_METRYKA[x][0], style='znaki-tabelki')
                komorka.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                komorka.add_paragraph()
                komorka.paragraphs[1].add_run(self.TEKSTY_METRYKA[x][1], style='znaki-tabelki')
                komorka.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
                if not jestem_pylem:
                    komorka.add_paragraph()
                    komorka.paragraphs[2].add_run(self.TEKSTY_METRYKA[x][2], style='znaki-tabelki')
                    komorka.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                komorka.paragraphs[0].add_run(self.TEKSTY_METRYKA[x], style='znaki-tabelki')
                komorka.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wierszy #6-8 prawej kolumny
        # wypełnienie wiersza #6 (WYNIK POMIARU)
        k6 = tabelka.columns[2].cells[5]
        k6.paragraphs[0].add_run(self.TEKSTY_WYNIK[0], style='znaki-tabelki')
        k6.paragraphs[0].add_run(self.TEKSTY_WYNIK[1], style='znaki-tabelki')
        k6.paragraphs[0].add_run(self.TEKSTY_WYNIK[2] + self.wdo, style='znaki-tabelki')
        k6.paragraphs[0].runs[1].font.superscript = True
        k6.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wiersza #7 (INTERPRETACJA WYNIKU)
        # obsługa wartości dopuszczalnej
        k7 = tabelka.columns[2].cells[6]
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP[0], style='znaki-tabelki')
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP[1], style='znaki-tabelki')
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP[2] + self.ndw, style='znaki-tabelki')
        k7.paragraphs[0].runs[1].font.superscript = True  # a niżej, dokładnie w takiej samej sytuacji musi być runs[2] - witaj python-docsie, żegnaj logiko!
        k7.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # obsługa oceny
        k7.add_paragraph()
        if u"<" in self.krotnosc or u">" in self.krotnosc:
            krot_do_oceny = float(self.krotnosc[1:])
        else:
            krot_do_oceny = float(self.krotnosc)
        if krot_do_oceny > 1.0:
            k7.paragraphs[1].add_run(self.TEKSTY_OCENA[1], style='znaki-tabelki')
        else:
            k7.paragraphs[1].add_run(self.TEKSTY_OCENA[0], style='znaki-tabelki')
        k7.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # obsługa krotności
        k7.add_paragraph()
        k7.paragraphs[2].add_run(self.TEKSTY_KROTNOSC + self.krotnosc, style='znaki-tabelki')
        k7.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # k7.paragraphs[1].paragraph_format.space_before = Pt(3)
        # k7.paragraphs[2].paragraph_format.space_before = Pt(3)

        # wypełnienie wiersza #9 (P.O.) - jeśli trzeba
        if self.ilosc_wierszy_tab == 9:
            k9 = tabelka.columns[2].cells[8]
            k9.paragraphs[0].add_run(self.TEKSTY_P_O[0] + self.p_o, style='znaki-tabelki')
            k9.paragraphs[0].add_run(self.TEKSTY_P_O[1], style='znaki-tabelki')
            k9.paragraphs[0].add_run(self.TEKSTY_P_O[2], style='znaki-tabelki')
            k9.paragraphs[0].add_run(self.TEKSTY_P_O[3], style='znaki-tabelki')
            k9.paragraphs[0].runs[2].font.superscript = True
            k9.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # stopka tabelki
        if not jestem_pylem:
            stopka = karta.add_paragraph(style='akapit-stopki')
            stopka.add_run(self.TESKTY_STOPKA, style='znaki-stopki')
            stopka.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # interlinia
        karta.add_paragraph(style='akapit-stopki')

        return tabelka


class PylSiO2(Pylochem):
    """Podtyp Czynnika II stopnia - pył z krzemionką"""
    # stałe klasy
    # uwaga na spacje na końcu stringów
    TEKSTY_ZAW_SiO2 = u"Zawartość wolnej, krystalicznej krzemionki [%]: "

    TEKSTY_CALK_RESP = (u"(frakcja wdychalna)", u"(frakcja respirabilna)", u"pyłu frakcji wdychalnej", u"pyłu frakcji respirabilnej", u"frakcja respirabilna – nie podlega ocenie")
    # '3' musi być w superscripcie (Font object, superscript property, value: True) - dlatego tekst podzielony na 3 części
    TEKSTY_WYNIK_SiO2 = (u"Wskaźniki narażenia [mg/m", u"3", u"]:")

    TEKSTY_WARTOSCI_DOP_SiO2 = (u"Wartości dopuszczalne wskaźników narażenia [mg/m", u"3", u"]:")

    TEKSTY_OCENA_SiO2 = (u"Wskaźnik narażenia ", u" nie przekracza wartości dopuszczalnej.", u" przekracza wartość dopuszczalną.")

    TEKSTY_KROTNOSC_SiO2 = (u"Krotność wartości dopuszczalnej wskaźnika narażenia ", u": ")

    TEKSTY_P_O_SiO2 = (u"Stężenie ", u" poniżej oznaczalności – poniżej ", u" [mg/m", u"3", u"]")

    def __init__(self, nazwa, nr_spr, miejsce, nazwa_st, data, nds, wskaznik_nar, krotnosc, czynnosci, wykonawca, metoda, numer, p_o, zaw_SiO2, nds_resp, wskaznik_nar_resp, krotnosc_resp, p_o_resp):
        super(PylSiO2, self).__init__(nazwa, nr_spr, miejsce, nazwa_st, data, nds, wskaznik_nar, krotnosc, czynnosci, wykonawca, metoda, numer, p_o)
        self.zaw_SiO2 = zaw_SiO2
        self.nds_resp = nds_resp  # NDS pyłu respirabilnego
        self.wskaznik_nar_resp = wskaznik_nar_resp  # wskaźnik narażenia pyły respirabilnego
        self.krotnosc_resp = krotnosc_resp
        self.p_o_resp = p_o_resp

    # TODO: space_before i space_after nie działają prawidłowo (wina python-docx?)
    def rysujTabelke(self, karta):
        """
        Rysuje tabelkę tego Czynnika. Większość rysowania oddaje do tej samej metody w klasie rodzica
        """
        # całą tabelkę, nagłówek i stopkę wypełnia rodzic, ta metoda zmienia tylko 2 (lub 3 ze wzgl. p_o) wiersze prawej kolumny
        tabelka = super(PylSiO2, self).rysujTabelke(karta)

        # wypełnienie wierszy #6-8 prawej kolumny
        # wypełnienie wiersza #6 (WYNIK POMIARU)
        k6 = tabelka.columns[2].cells[5]
        # UWAGA! najpierw wyczyszczenie całej komórki
        k6.text = u""
        k6.paragraphs[0].add_run(self.TEKSTY_ZAW_SiO2 + self.zaw_SiO2, style='znaki-tabelki')
        k6.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        k6.add_paragraph()
        k6.paragraphs[1].add_run(self.TEKSTY_WYNIK_SiO2[0], style='znaki-tabelki')
        k6.paragraphs[1].add_run(self.TEKSTY_WYNIK_SiO2[1], style='znaki-tabelki')
        k6.paragraphs[1].add_run(self.TEKSTY_WYNIK_SiO2[2], style='znaki-tabelki')
        k6.paragraphs[1].runs[1].font.superscript = True
        k6.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k6.add_paragraph()
        k6.paragraphs[2].add_run(self.wdo + u" " + self.TEKSTY_CALK_RESP[0], style='znaki-tabelki')
        k6.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k6.add_paragraph()
        k6.paragraphs[3].add_run(self.wskaznik_nar_resp + u" " + self.TEKSTY_CALK_RESP[1], style='znaki-tabelki')
        k6.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k6.paragraphs[1].paragraph_format.space_before = Pt(1.5)

        # wypełnienie wiersza #7 (INTERPRETACJA WYNIKU)
        # obsługa wartości dopuszczalnej
        k7 = tabelka.columns[2].cells[6]
        # UWAGA! najpierw wyczyszczenie całej komórki
        k7.text = u""
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP_SiO2[0], style='znaki-tabelki')
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP_SiO2[1], style='znaki-tabelki')
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP_SiO2[2], style='znaki-tabelki')
        k7.paragraphs[0].runs[2].font.superscript = True  # tu musi być run[2] mimo że niżej taka sama sytuacja a musi być run[1] - witaj python-docsie, żegnaj logiko!
        k7.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        k7.add_paragraph()
        k7.paragraphs[1].add_run(self.ndw + u" " + self.TEKSTY_CALK_RESP[0], style='znaki-tabelki')
        k7.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k7.add_paragraph()
        if self.nds_resp == u"brak":
            k7.paragraphs[2].add_run(self.TEKSTY_CALK_RESP[4], style='znaki-tabelki')
        else:
            k7.paragraphs[2].add_run(self.nds_resp + u" " + self.TEKSTY_CALK_RESP[1], style='znaki-tabelki')
        k7.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # obsługa oceny
        k7.add_paragraph()
        if u"<" in self.krotnosc or u">" in self.krotnosc:
            krot_do_oceny = float(self.krotnosc[1:])
        else:
            krot_do_oceny = float(self.krotnosc)
        if krot_do_oceny > 1.0:
            k7.paragraphs[3].add_run(self.TEKSTY_OCENA_SiO2[0] + self.TEKSTY_CALK_RESP[2] + self.TEKSTY_OCENA_SiO2[2], style='znaki-tabelki')
        else:
            k7.paragraphs[3].add_run(self.TEKSTY_OCENA_SiO2[0] + self.TEKSTY_CALK_RESP[2] + self.TEKSTY_OCENA_SiO2[1], style='znaki-tabelki')
        k7.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k7.paragraphs[2].paragraph_format.space_before = Pt(1)

        dodane_akapity = 0
        if self.krotnosc_resp != u"-":
            k7.add_paragraph()
            if u"<" in self.krotnosc_resp or u">" in self.krotnosc_resp:
                krot_resp_do_oceny = float(self.krotnosc_resp[1:])
            else:
                krot_resp_do_oceny = float(self.krotnosc_resp)
            if krot_resp_do_oceny > 1.0:
                k7.paragraphs[4].add_run(self.TEKSTY_OCENA_SiO2[0] + self.TEKSTY_CALK_RESP[3] + self.TEKSTY_OCENA_SiO2[2], style='znaki-tabelki')
            else:
                k7.paragraphs[4].add_run(self.TEKSTY_OCENA_SiO2[0] + self.TEKSTY_CALK_RESP[3] + self.TEKSTY_OCENA_SiO2[1], style='znaki-tabelki')
            k7.paragraphs[4].alignment = WD_ALIGN_PARAGRAPH.LEFT
            dodane_akapity += 1

        # obsługa krotności
        k7.add_paragraph()
        k7.paragraphs[4 + dodane_akapity].add_run(self.TEKSTY_KROTNOSC_SiO2[0] + self.TEKSTY_CALK_RESP[0] + self.TEKSTY_KROTNOSC_SiO2[1] + self.krotnosc, style='znaki-tabelki')
        k7.paragraphs[4 + dodane_akapity].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # k7.paragraphs[4 + dodane_akapity].paragraph_format.space_before = Pt(3)

        if self.krotnosc_resp != u"-":
            k7.add_paragraph()
            k7.paragraphs[6].add_run(self.TEKSTY_KROTNOSC_SiO2[0] + self.TEKSTY_CALK_RESP[1] + self.TEKSTY_KROTNOSC_SiO2[1] + self.krotnosc_resp, style='znaki-tabelki')
            k7.paragraphs[6].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # TODO: sprawdzić działanie superscriptu w Calkowitym - musiałem zmienić index runa w linii #477 na 2gi, żeby działało prawidłowo, podczas gdy w Respirabilnym jest index 1szy i nie widzę różnicy - o co chodzi?
        # wypełnienie wiersza #9 (P.O.) - jeśli trzeba
        def dodajAkapitPOCalk(akapit, teksty_p_o_SiO2, teksty_calk_resp, p_o):
            """Dodaje do komórki odpowiednio zformatowany akapit z p.o. pyłu całkowitego"""
            akapit.add_run(teksty_p_o_SiO2[0] + teksty_calk_resp[2] + teksty_p_o_SiO2[1] + p_o + teksty_p_o_SiO2[2], style='znaki-tabelki')
            akapit.add_run(teksty_p_o_SiO2[3], style='znaki-tabelki')
            akapit.add_run(teksty_p_o_SiO2[4], style='znaki-tabelki')
            ilosc_linijek = len(akapit.runs)
            akapit.runs[ilosc_linijek - 2].font.superscript = True
            akapit.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def dodajAkapitPOResp(akapit, teksty_p_o_SiO2, teksty_calk_resp, p_o_resp):
            """Dodaje do komórki odpowiednio zformatowany akapit z p.o. pyłu respirabilnego"""
            akapit.add_run(teksty_p_o_SiO2[0] + teksty_calk_resp[3] + teksty_p_o_SiO2[1] + p_o_resp + teksty_p_o_SiO2[2], style='znaki-tabelki')
            akapit.add_run(teksty_p_o_SiO2[3], style='znaki-tabelki')
            akapit.add_run(teksty_p_o_SiO2[4], style='znaki-tabelki')
            ilosc_linijek = len(akapit.runs)
            akapit.runs[ilosc_linijek - 2].font.superscript = True
            akapit.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if self.ilosc_wierszy_tab == 9 and self.p_o_resp is not None:
            k9 = tabelka.columns[2].cells[8]
            # UWAGA! najpierw wyczyszczenie całej komórki
            k9.text = u""
            dodajAkapitPOCalk(k9.paragraphs[0], self.TEKSTY_P_O_SiO2, self.TEKSTY_CALK_RESP, self.p_o)
            k9.add_paragraph()
            dodajAkapitPOResp(k9.paragraphs[1], self.TEKSTY_P_O_SiO2, self.TEKSTY_CALK_RESP, self.p_o_resp)

        elif self.ilosc_wierszy_tab == 9 and self.p_o_resp is None:
            k9 = tabelka.columns[2].cells[8]
            # UWAGA! najpierw wyczyszczenie całej komórki
            k9.text = u""
            dodajAkapitPOCalk(k9.paragraphs[0], self.TEKSTY_P_O_SiO2, self.TEKSTY_CALK_RESP, self.p_o)

        elif self.ilosc_wierszy_tab == 8 and self.p_o_resp is not None:
            self.ilosc_wierszy_tab == 9
            ost_wiersz = tabelka.add_row()
            komorka_lewa = ost_wiersz.cells[0]
            komorka_lewa.paragraphs[0].add_run(u"9.", style='znaki-tabelki')
            komorka_lewa.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            komorka_srodkowa = ost_wiersz.cells[1]
            komorka_srodkowa.paragraphs[0].add_run(u"Uwagi", style='znaki-tabelki')
            komorka_srodkowa.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            komorka_prawa = ost_wiersz.cells[2]
            dodajAkapitPOResp(komorka_prawa.paragraphs[0], self.TEKSTY_P_O_SiO2, self.TEKSTY_CALK_RESP, self.p_o_resp)


class ChemiaChwilowkowa(Pylochem):
    """Podtyp Czynnika II stopnia - chemia z chwilówkami"""
    # uwaga na spacje na końcu stringów
    # '3' musi być w superscripcie (Font object, superscript property, value: True) - dlatego tekst podzielony na 3 części
    TEKSTY_WYNIK_CHW = (u"Najwyższe stężenie chwilowe [mg/m", u"3", u"]: ")

    TEKSTY_WARTOSCI_DOP_CHW = (u"Wartość dopuszczalna stężenia chwilowego [mg/m", u"3", u"]: ")

    TEKSTY_OCENA_CHW = (u"Najwyższe stężenie chwilowe nie przekracza wartości dopuszczalnej.", u"Najwyższe stężenie chwilowe przekracza wartość dopuszczalną.")

    TEKSTY_KROTNOSC_CHW = u"Krotność wartości dopuszczalnej najwyższego stężenia chwilowego: "

    TEKSTY_P_O_CHW = (u"Najwyższe stężenie chwilowe poniżej oznaczalności – poniżej ", u" [mg/m", u"3", u"]")

    def __init__(self, nazwa, nr_spr, miejsce, nazwa_st, data, nds, wskaznik_nar, krotnosc, czynnosci, wykonawca, metoda, numer, p_o, nds_chw, wskaznik_nar_chw, krotnosc_chw, p_o_chw):
        super(ChemiaChwilowkowa, self).__init__(nazwa, nr_spr, miejsce, nazwa_st, data, nds, wskaznik_nar, krotnosc, czynnosci, wykonawca, metoda, numer, p_o)
        self.nds_chw = nds_chw
        self.wskaznik_nar_chw = wskaznik_nar_chw
        self.krotnosc_chw = krotnosc_chw
        self.p_o_chw = p_o_chw

    # TODO: space_before i space_after nie działają prawidłowo (wina python-docx?)
    def rysujTabelke(self, karta):
        """
        Rysuje tabelkę tego Czynnika. Większość rysowania oddaje do tej samej metody w klasie rodzica
        """
        # całą tabelkę, nagłówek i stopkę wypełnia rodzic, ta metoda zmienia tylko 2 (lub 3 ze wzgl. p_o) wiersze prawej kolumny
        tabelka = super(ChemiaChwilowkowa, self).rysujTabelke(karta)

        # wypełnienie wierszy #6-8 prawej kolumny
        # wypełnienie wiersza #6 (WYNIK POMIARU)
        k6 = tabelka.columns[2].cells[5]
        k6.add_paragraph()
        k6.paragraphs[1].add_run(self.TEKSTY_WYNIK_CHW[0], style='znaki-tabelki')
        k6.paragraphs[1].add_run(self.TEKSTY_WYNIK_CHW[1], style='znaki-tabelki')
        k6.paragraphs[1].add_run(self.TEKSTY_WYNIK_CHW[2] + self.wskaznik_nar_chw, style='znaki-tabelki')
        k6.paragraphs[1].runs[1].font.superscript = True
        k6.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # wypełnienie wiersza #7 (INTERPRETACJA WYNIKU)
        # obsługa wartości dopuszczalnej
        k7 = tabelka.columns[2].cells[6]
        # UWAGA! najpierw wyczyszczenie całej komórki
        k7.text = u""
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP[0], style='znaki-tabelki')
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP[1], style='znaki-tabelki')
        k7.paragraphs[0].add_run(self.TEKSTY_WARTOSCI_DOP[2] + self.ndw, style='znaki-tabelki')
        k7.paragraphs[0].runs[2].font.superscript = True  # tu musi być run[2] mimo że niżej taka sama sytuacja a musi być run[1] - witaj python-docsie, żegnaj logiko!
        k7.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k7.add_paragraph()
        k7.paragraphs[1].add_run(self.TEKSTY_WARTOSCI_DOP_CHW[0], style='znaki-tabelki')
        k7.paragraphs[1].add_run(self.TEKSTY_WARTOSCI_DOP_CHW[1], style='znaki-tabelki')
        k7.paragraphs[1].add_run(self.TEKSTY_WARTOSCI_DOP_CHW[2] + self.nds_chw, style='znaki-tabelki')
        k7.paragraphs[1].runs[1].font.superscript = True
        k7.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # obsługa oceny
        k7.add_paragraph()
        if u"<" in self.krotnosc or u">" in self.krotnosc:
            krot_do_oceny = float(self.krotnosc[1:])
        else:
            krot_do_oceny = float(self.krotnosc)
        if krot_do_oceny > 1.0:
            k7.paragraphs[2].add_run(self.TEKSTY_OCENA[1], style='znaki-tabelki')
        else:
            k7.paragraphs[2].add_run(self.TEKSTY_OCENA[0], style='znaki-tabelki')
        k7.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k7.add_paragraph()
        if u"<" in self.krotnosc_chw or u">" in self.krotnosc_chw:
            krot_chw_do_oceny = float(self.krotnosc_chw[1:])
        else:
            krot_chw_do_oceny = float(self.krotnosc_chw)
        if krot_chw_do_oceny > 1.0:
            k7.paragraphs[3].add_run(self.TEKSTY_OCENA_CHW[1], style='znaki-tabelki')
        else:
            k7.paragraphs[3].add_run(self.TEKSTY_OCENA_CHW[0], style='znaki-tabelki')
        k7.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # obsługa krotności
        k7.add_paragraph()
        k7.paragraphs[4].add_run(self.TEKSTY_KROTNOSC + self.krotnosc, style='znaki-tabelki')
        k7.paragraphs[4].alignment = WD_ALIGN_PARAGRAPH.LEFT
        k7.add_paragraph()
        k7.paragraphs[5].add_run(self.TEKSTY_KROTNOSC_CHW + self.krotnosc_chw, style='znaki-tabelki')
        k7.paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # k7.paragraphs[2].paragraph_format.space_before = Pt(3)
        # k7.paragraphs[4].paragraph_format.space_before = Pt(3)

        def dodajAkapitPOChw(akapit, teksty_p_o_chw, p_o_chw):
            """Dodaje do komórki odpowiednio zformatowany akapit z p.o. próbki chwilowej"""
            akapit.add_run(teksty_p_o_chw[0] + p_o_chw, style='znaki-tabelki')
            akapit.add_run(teksty_p_o_chw[1], style='znaki-tabelki')
            akapit.add_run(teksty_p_o_chw[2], style='znaki-tabelki')
            akapit.add_run(teksty_p_o_chw[3], style='znaki-tabelki')
            ilosc_linijek = len(akapit.runs)
            akapit.runs[ilosc_linijek - 2].font.superscript = True
            akapit.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # # wypełnienie wiersza #9 (P.O.) - jeśli trzeba
        if self.ilosc_wierszy_tab == 9 and self.p_o_chw is not None:
            k9 = tabelka.columns[2].cells[8]
            k9.add_paragraph()
            dodajAkapitPOChw(k9.paragraphs[1], self.TEKSTY_P_O_CHW, self.p_o_chw)

        elif self.ilosc_wierszy_tab == 8 and self.p_o_chw is not None:
            self.ilosc_wierszy_tab == 9
            ost_wiersz = tabelka.add_row()
            komorka_lewa = ost_wiersz.cells[0]
            komorka_lewa.paragraphs[0].add_run(u"9.", style='znaki-tabelki')
            komorka_lewa.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            komorka_srodkowa = ost_wiersz.cells[1]
            komorka_srodkowa.paragraphs[0].add_run(u"Uwagi", style='znaki-tabelki')
            komorka_srodkowa.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            komorka_prawa = ost_wiersz.cells[2]
            dodajAkapitPOChw(komorka_prawa.paragraphs[0], self.TEKSTY_P_O_CHW, self.p_o_chw)
